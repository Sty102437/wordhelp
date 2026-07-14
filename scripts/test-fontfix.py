#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Test script for fix-cjk-fonts.ps1 font fix logic.

Creates a test .docx with mixed CJK/EN runs in paragraphs, tables,
and headers, then verifies that fix-cjk-fonts correctly sets rFonts.

Usage: python test-fontfix.py
"""
import sys
import os
import tempfile

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def create_test_docx(path):
    """Create a test document with mixed CJK/EN content."""
    doc = Document()

    # Paragraph with mixed content
    p = doc.add_paragraph("图1-3 系统架构图")
    p2 = doc.add_paragraph("This is a test 测试中英文混排")

    # Table with mixed content
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "项目名称"
    table.cell(0, 1).text = "Project Name 项目名称"

    # Header with mixed content
    section = doc.sections[0]
    section.header.paragraphs[0].text = "报告标题 Report Title"

    doc.save(path)
    return path


def verify_fix(path):
    """Verify that rFonts are correctly set after fix-cjk-fonts runs."""
    doc = Document(path)

    passed = 0
    failed = 0

    def check_run(run, location):
        nonlocal passed, failed
        text = run.text.strip()
        if not text:
            return

        has_cn = bool(__import__('re').search(r'[\u4e00-\u9fff]', text))
        has_en = bool(__import__('re').search(r'[a-zA-Z0-9]', text))

        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            print(f"  FAIL: {location} — no rPr found")
            failed += 1
            return

        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            print(f"  FAIL: {location} — no rFonts found")
            failed += 1
            return

        if has_cn:
            east_asia = rFonts.get(qn('w:eastAsia'))
            if east_asia == 'SimSun':
                print(f"  OK:   {location} — eastAsia=SimSun")
                passed += 1
            else:
                print(f"  FAIL: {location} — eastAsia='{east_asia}' (expected SimSun)")
                failed += 1

        if has_en:
            ascii_font = rFonts.get(qn('w:ascii'))
            if ascii_font == 'Times New Roman':
                print(f"  OK:   {location} — ascii=Times New Roman")
                passed += 1
            else:
                print(f"  FAIL: {location} — ascii='{ascii_font}' (expected Times New Roman)")
                failed += 1

    # Check paragraphs
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            check_run(run, f"P{i}")

    # Check tables
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, para in enumerate(cell.paragraphs):
                    for run in para.runs:
                        check_run(run, f"T{ti}R{ri}C{ci}P{pi}")

    # Check headers
    for si, section in enumerate(doc.sections):
        for pi, para in enumerate(section.header.paragraphs):
            for run in para.runs:
                check_run(run, f"HeaderS{si}P{pi}")

    print(f"\nResults: {passed} passed, {failed} failed")
    return failed == 0


if __name__ == "__main__":
    test_doc = os.path.join(tempfile.gettempdir(), "test-fontfix-input.docx")
    output_doc = os.path.join(tempfile.gettempdir(), "test-fontfix-output.docx")

    print("=== wordhelp Font Fix Test ===\n")

    # Step 1: Create test document
    create_test_docx(test_doc)
    print(f"Created test document: {test_doc}\n")

    # Step 2: Run fix-cjk-fonts via PowerShell
    import subprocess
    result = subprocess.run(
        ["powershell", "-File", "fix-cjk-fonts.ps1",
         "-InputPath", test_doc, "-OutputPath", output_doc],
        capture_output=True, text=True,
        cwd=os.path.dirname(os.path.abspath(__file__))
    )
    print(result.stdout)
    if result.stderr:
        print(f"PowerShell errors: {result.stderr}")

    # Step 3: Verify
    if os.path.exists(output_doc):
        print("\n--- Verification ---")
        success = verify_fix(output_doc)
        if success:
            print("\nAll tests passed!")
        else:
            print("\nSome tests failed!")
            sys.exit(1)
    else:
        print(f"ERROR: Output file not created: {output_doc}")
        sys.exit(1)
