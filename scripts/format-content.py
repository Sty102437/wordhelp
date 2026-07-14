#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Format document content using two-pass approach: insert first, format second.

Pass 1: Content is already inserted (this script assumes content is in the doc).
Pass 2: Apply heading styles and font formatting based on text patterns.

Usage: python format-content.py --input template.docx --output output.docx [--type academic]
"""
import sys
import re
import argparse

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# --- Pattern matchers ---
def is_heading1(text):
    return bool(re.match(r'^\d+(\.\d+)+\s', text)) or text.startswith('第') and '章' in text[:10]

def is_heading2(text):
    return bool(re.match(r'^\d+\.\d+\s', text))

def is_heading3(text):
    return bool(re.match(r'^\d+\.\d+\.\d+\s', text))

def is_abstract_title(text):
    return text in ('摘要', 'Abstract', '摘 要')

def is_references_title(text):
    return text in ('参考文献', 'References', '参 考 文 献')

def is_keywords(text):
    return text.startswith('关键词') or text.startswith('Keywords')

def is_reference_item(text):
    return bool(re.match(r'^\[\d+\]', text))


def set_font(run, cn_font='宋体', en_font='Times New Roman', size=None, bold=None):
    """Set font on a run using rFonts to avoid overwrite issues."""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def set_outline_level(para, level):
    """Set outline level without applying Heading style (for Abstract/References)."""
    pPr = para._element.get_or_add_pPr()
    outline = pPr.find(qn('w:outlineLvl'))
    if outline is None:
        outline = OxmlElement('w:outlineLvl')
        pPr.append(outline)
    outline.set(qn('w:val'), str(level))


def format_document(doc, doc_type='academic'):
    """Pass 2: Apply formatting based on text patterns."""
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if is_heading1(text):
            para.style = doc.styles['Heading 1']
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                set_font(run, size=16, bold=True)

        elif is_heading2(text):
            para.style = doc.styles['Heading 2']
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                set_font(run, size=15, bold=True)

        elif is_heading3(text):
            para.style = doc.styles['Heading 3']
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                set_font(run, size=14, bold=False)

        elif is_abstract_title(text):
            # NO heading style — use outlineLvl to appear in TOC without numbering
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_outline_level(para, 0)
            for run in para.runs:
                set_font(run, size=16, bold=True)

        elif is_references_title(text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_outline_level(para, 0)
            for run in para.runs:
                set_font(run, size=16, bold=True)

        elif is_keywords(text):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                set_font(run, size=12, bold=False)
            # Bold the "关键词:" / "Keywords:" prefix
            if para.runs:
                para.runs[0].font.bold = True

        elif is_reference_item(text):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                set_font(run, size=10.5, bold=False)

        else:
            # Body text
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                set_font(run, size=12, bold=False)

    # Equation numbering tab stops (if any paragraph looks like an equation)
    for para in doc.paragraphs:
        if re.search(r'\(\d+-\d+\)\s*$', para.text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.tab_stops.add_tab_stop(Cm(14.66), WD_TAB_ALIGNMENT.RIGHT)


def mark_toc_update(doc):
    """Mark TOC fields for update when Word opens the document."""
    for rel in doc.part.rels.values():
        if "TOC" in str(rel.reltype):
            pass  # python-docx doesn't directly expose TOC fields
    # Set updateFields in settings
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings.append(update_fields)
    update_fields.set(qn('w:val'), 'true')


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Format .docx content (Pass 2)")
    parser.add_argument("--input", required=True, help="Path to input .docx (content already inserted)")
    parser.add_argument("--output", required=True, help="Path to output .docx")
    parser.add_argument("--type", default="academic", help="Document type: academic | general | government")
    args = parser.parse_args()

    doc = Document(args.input)
    format_document(doc, args.type)
    mark_toc_update(doc)
    doc.save(args.output)
    print(f"Formatted: {args.output}")
