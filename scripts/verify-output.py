#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Verify output .docx structure: count paragraphs, tables, check styles.

Usage: python verify-output.py --input output.docx
"""
import sys
import argparse

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def verify(input_path):
    doc = Document(input_path)

    print(f"=== Output Verification: {input_path} ===\n")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    print(f"Sections: {len(doc.sections)}")

    # Count images
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in str(rel.reltype):
            image_count += 1
    print(f"Images: {image_count}")

    # First 10 non-empty paragraphs with styles
    print("\n--- First 10 paragraphs ---")
    shown = 0
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"  P{i} [{p.style.name}] {p.text[:80]}")
            shown += 1
            if shown >= 10:
                break

    # Tables summary
    if doc.tables:
        print(f"\n--- {len(doc.tables)} table(s) ---")
        for i, t in enumerate(doc.tables):
            first_cell = t.cell(0, 0).text[:50] if t.rows else "(empty)"
            print(f"  T{i}: {len(t.rows)}r x {len(t.columns)}c — {first_cell}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Verify .docx output structure")
    parser.add_argument("--input", required=True, help="Path to .docx file")
    args = parser.parse_args()
    verify(args.input)
