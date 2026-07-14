#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Analyze a .docx template to identify cover page, scoring page, TOC, and body sections.

Usage: python analyze-template.py --input template.docx
"""
import sys
import argparse

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def analyze(input_path):
    doc = Document(input_path)

    print(f"=== Template Analysis: {input_path} ===\n")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    print(f"Total sections: {len(doc.sections)}\n")

    # Paragraphs
    print("--- Paragraphs ---")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"P{i} [{p.style.name}] {p.text[:80]}")

    # Tables
    print("\n--- Tables ---")
    for i, t in enumerate(doc.tables):
        first_cell = t.cell(0, 0).text[:50] if t.rows else "(empty)"
        print(f"T{i}: {len(t.rows)}r x {len(t.columns)}c, first cell: {first_cell}")

    # Section hints
    print("\n--- Section Detection Hints ---")
    keywords_scoring = ["评分", "score", "考核", "评分标准", "scoring standard"]
    keywords_cover = ["封面", "cover", "题目", "姓名", "学号", "日期"]

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().lower()
        for kw in keywords_scoring:
            if kw in text:
                print(f"  [SCORING?] P{i}: {p.text[:80]}")
        for kw in keywords_cover:
            if kw in text and i < 10:
                print(f"  [COVER?]   P{i}: {p.text[:80]}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Analyze .docx template structure")
    parser.add_argument("--input", required=True, help="Path to .docx file")
    args = parser.parse_args()
    analyze(args.input)
