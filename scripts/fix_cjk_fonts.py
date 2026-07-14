#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Fix CJK/EN fonts at the run level — cross-platform.

Sets w:rFonts eastAsia for Chinese characters, ascii/hAnsi for English/numbers.
Processes body paragraphs, table cells, headers, and footers.

Usage: python fix_cjk_fonts.py --input output.docx [--output fixed.docx]
                                  [--cn-font 宋体] [--en-font "Times New Roman"]
"""
import sys
import os
import re
import argparse

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


cn_pattern = re.compile(r'[\u4e00-\u9fff\u3400-\u4dbf]')
ascii_pattern = re.compile(r'[a-zA-Z0-9]')


def fix_run(run, cn_font, en_font):
    """Fix font on a single run using w:rFonts attributes."""
    if not run.text.strip():
        return

    has_cn = bool(cn_pattern.search(run.text))
    has_en = bool(ascii_pattern.search(run.text))

    if not has_cn and not has_en:
        return

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    if has_cn:
        rFonts.set(qn('w:eastAsia'), cn_font)
    if has_en:
        rFonts.set(qn('w:ascii'), en_font)
        rFonts.set(qn('w:hAnsi'), en_font)


def fix_paragraphs(paragraphs, cn_font, en_font):
    """Fix all runs in a list of paragraphs."""
    for para in paragraphs:
        for run in para.runs:
            fix_run(run, cn_font, en_font)


def fix_document(input_path, output_path, cn_font, en_font):
    """Fix CJK/EN fonts in a .docx file."""
    doc = Document(input_path)

    # Body paragraphs
    fix_paragraphs(doc.paragraphs, cn_font, en_font)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fix_paragraphs(cell.paragraphs, cn_font, en_font)

    # Headers and footers
    for section in doc.sections:
        fix_paragraphs(section.header.paragraphs, cn_font, en_font)
        fix_paragraphs(section.footer.paragraphs, cn_font, en_font)

    doc.save(output_path)
    print(f"Fixed: {output_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Fix CJK/EN fonts in .docx (cross-platform)")
    parser.add_argument("--input", "-i", required=True, help="Input .docx file")
    parser.add_argument("--output", "-o", help="Output .docx (default: <name>.fixed.docx)")
    parser.add_argument("--cn-font", default="SimSun", help="Chinese font (default: SimSun)")
    parser.add_argument("--en-font", default="Times New Roman", help="English font (default: Times New Roman)")
    args = parser.parse_args()

    if not args.output:
        base = os.path.splitext(args.input)[0]
        ext = os.path.splitext(args.input)[1]
        args.output = f"{base}.fixed{ext}"

    fix_document(args.input, args.output, args.cn_font, args.en_font)
